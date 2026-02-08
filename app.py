"""
Сервис расчёта КП с автоматическим ценообразованием
Раздельные таблицы для РБ и ФБ контрактов
"""

import streamlit as st
import pandas as pd
from io import BytesIO
import sys
import os
import pickle
import json
import base64
from pathlib import Path

import requests as http_requests
from dotenv import load_dotenv

load_dotenv()

# Добавляем путь к модулям
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from src.parsers.cost_parser import parse_cost_file
from src.parsers.competitor_parser import parse_competitor_file
from src.parsers.request_parser import parse_request_file
from src.matching.product_matcher import match_products
from src.calculator.pricing import calculate_prices
from src.calculator.economics import calculate_economics
from src.export.excel_export import export_kp_to_excel, export_economics_to_excel
from src.export.docx_export import export_kp_to_docx

# Директория для сохранения данных
CACHE_DIR = Path(__file__).parent / ".cache"
CACHE_DIR.mkdir(exist_ok=True)

def save_data(key: str, data):
    """Сохранить данные на диск"""
    try:
        with open(CACHE_DIR / f"{key}.pkl", "wb") as f:
            pickle.dump(data, f)
    except Exception as e:
        print(f"Ошибка сохранения {key}: {e}")

def load_data(key: str):
    """Загрузить данные с диска"""
    try:
        cache_file = CACHE_DIR / f"{key}.pkl"
        if cache_file.exists():
            with open(cache_file, "rb") as f:
                return pickle.load(f)
    except Exception as e:
        print(f"Ошибка загрузки {key}: {e}")
    return None

# Настройка страницы
st.set_page_config(
    page_title="Расчёт КП",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Стили
st.markdown("""
<style>
    .main-header {
        font-size: 1.8rem;
        font-weight: bold;
        text-align: center;
        padding: 1.2rem 1rem;
        background: linear-gradient(135deg, #1e3a5f, #2e5a8f);
        color: white;
        border-radius: 10px;
        margin: 0.5rem 0 1.5rem 0;
        line-height: 1.4;
    }
    .contract-header {
        font-size: 1.2rem;
        font-weight: bold;
        padding: 0.5rem;
        background: #e8eef5;
        color: #1e3a5f !important;
        border-radius: 5px;
        text-align: center;
        margin-bottom: 0.5rem;
    }
    .contract-header-rb {
        font-size: 1.2rem;
        font-weight: bold;
        padding: 0.5rem;
        background: #dbeafe;
        color: #1e40af !important;
        border-radius: 5px;
        text-align: center;
        margin-bottom: 0.5rem;
    }
    .contract-header-fb {
        font-size: 1.2rem;
        font-weight: bold;
        padding: 0.5rem;
        background: #dcfce7;
        color: #166534 !important;
        border-radius: 5px;
        text-align: center;
        margin-bottom: 0.5rem;
    }
    .stDataFrame, [data-testid="stDataFrame"] {
        width: 100% !important;
    }
    [data-testid="stDataFrame"] > div {
        width: 100% !important;
    }
    [data-testid="stDataFrame"] td, [data-testid="stDataFrame"] th {
        color: #262730 !important;
    }
    .block-container {
        padding-top: 1rem !important;
        padding-bottom: 1rem !important;
    }
    .metrics-row {
        display: flex;
        gap: 10px;
        margin: 0.4rem 0 0.8rem 0;
    }
    .metric-card {
        flex: 1;
        padding: 10px 14px;
        border-radius: 8px;
        border: 1px solid #e0e0e0;
        background: #ffffff;
    }
    .metric-label {
        font-size: 0.72rem;
        font-weight: 600;
        color: #666 !important;
        margin-bottom: 4px;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    .metric-value {
        font-size: 1.25rem;
        font-weight: 700;
    }
    .metric-value.blue { color: #1565C0 !important; }
    .metric-value.green { color: #2E7D32 !important; }
    .metric-value.orange { color: #E65100 !important; }
    .metric-value.red { color: #C62828 !important; }
    .metric-value.teal { color: #00695C !important; }
    .metric-value.gray { color: #546E7A !important; }
    .summary-row {
        display: flex;
        gap: 10px;
        margin: 0.3rem 0;
    }
    .summary-card {
        flex: 1;
        padding: 12px 16px;
        border-radius: 8px;
        border: 1px solid #e0e0e0;
        background: #fafafa;
    }
    .summary-card.rb { border-left: 4px solid #1565C0; }
    .summary-card.fb { border-left: 4px solid #2E7D32; }
    .summary-card.total { border-left: 4px solid #00695C; background: #f0fdf4; }
    .summary-title {
        font-size: 0.8rem;
        font-weight: 700;
        color: #333 !important;
        margin-bottom: 6px;
    }
    .summary-line {
        font-size: 0.85rem;
        color: #444 !important;
        line-height: 1.6;
    }
    .summary-line b { font-weight: 700; }
</style>
""", unsafe_allow_html=True)

# Заголовок
st.markdown('<div class="main-header">📊 Сервис расчёта коммерческих предложений</div>', unsafe_allow_html=True)

# Инициализация состояния — загрузка с диска
if 'initialized' not in st.session_state:
    st.session_state.initialized = True
    st.session_state.cost_data = load_data('cost_data')
    st.session_state.competitor_data = load_data('competitor_data')
    st.session_state.rb_request = load_data('rb_request')
    st.session_state.fb_request = load_data('fb_request')
    st.session_state.rb_data = load_data('rb_data')
    st.session_state.fb_data = load_data('fb_data')
    st.session_state.loaded_files = load_data('loaded_files') or {'cost': None, 'competitor': None, 'rb': None, 'fb': None}

    # Миграция со старого формата (единая таблица → раздельные)
    if st.session_state.rb_data is None and st.session_state.fb_data is None:
        old = load_data('edited') or load_data('calculated')
        if old is not None and 'Контракт' in old.columns:
            rb = old[old['Контракт'] == 'РБ'].copy()
            fb = old[old['Контракт'] == 'ФБ'].copy()
            if len(rb) > 0:
                rb['№'] = range(1, len(rb) + 1)
                st.session_state.rb_data = rb
                save_data('rb_data', rb)
            if len(fb) > 0:
                fb['№'] = range(1, len(fb) + 1)
                st.session_state.fb_data = fb
                save_data('fb_data', fb)

# ============ ШАГ 1: Загрузка общих документов ============
col_step1_title, col_step1_clear = st.columns([6, 1])
with col_step1_title:
    st.subheader("📁 Шаг 1: Загрузка документов")
with col_step1_clear:
    if st.button("🗑️ Очистить кэш", help="Удалить все сохранённые данные и начать заново"):
        for f in CACHE_DIR.glob("*.pkl"):
            f.unlink()
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()

if st.session_state.cost_data is not None or st.session_state.competitor_data is not None:
    st.caption("💾 Сохранённые данные загружены. Загрузите новый файл для замены.")

col_cost, col_competitor = st.columns(2)

with col_cost:
    cost_file = st.file_uploader(
        "Себестоимость (Excel)",
        type=['xlsx', 'xls'],
        key="cost_uploader",
        help="Файл с себестоимостью товаров (Сравнение цен.xlsx)"
    )

    col_status, col_btn = st.columns([3, 1])
    with col_status:
        if st.session_state.cost_data is not None:
            st.success(f"✅ {len(st.session_state.cost_data)} позиций")
    with col_btn:
        if st.session_state.cost_data is not None:
            if st.button("🗑️", key="clear_cost", help="Очистить"):
                st.session_state.cost_data = None
                st.session_state.loaded_files['cost'] = None
                save_data('cost_data', None)
                save_data('loaded_files', st.session_state.loaded_files)
                st.rerun()

    if cost_file:
        if st.session_state.loaded_files['cost'] != cost_file.name:
            try:
                parsed = parse_cost_file(cost_file)
                if parsed is None or len(parsed) == 0:
                    st.error("Файл себестоимости пуст или не содержит данных с ценами > 0")
                else:
                    st.session_state.cost_data = parsed
                    st.session_state.loaded_files['cost'] = cost_file.name
                    save_data('cost_data', st.session_state.cost_data)
                    save_data('loaded_files', st.session_state.loaded_files)
                    st.rerun()
            except Exception as e:
                st.error(f"Ошибка: {e}")

with col_competitor:
    competitor_file = st.file_uploader(
        "КП конкурента (Word)",
        type=['docx'],
        key="competitor_uploader",
        help="Файл с ценами конкурента (.docx)"
    )

    col_status, col_btn = st.columns([3, 1])
    with col_status:
        if st.session_state.competitor_data is not None:
            total = st.session_state.competitor_data['Сумма'].sum() if 'Сумма' in st.session_state.competitor_data.columns else 0
            st.success(f"✅ {len(st.session_state.competitor_data)} поз., {total:,.0f} ₽")
    with col_btn:
        if st.session_state.competitor_data is not None:
            if st.button("🔄", key="reparse_competitor", help="Перераспознать"):
                st.session_state.loaded_files['competitor'] = None
                save_data('loaded_files', st.session_state.loaded_files)
                st.rerun()

    if competitor_file:
        if st.session_state.loaded_files['competitor'] != competitor_file.name:
            try:
                with st.spinner("📄 Парсинг Word..."):
                    st.session_state.competitor_data = parse_competitor_file(competitor_file, competitor_file.name)
                st.session_state.loaded_files['competitor'] = competitor_file.name
                save_data('competitor_data', st.session_state.competitor_data)
                save_data('loaded_files', st.session_state.loaded_files)
                st.rerun()
            except Exception as e:
                st.error(f"Ошибка: {e}")

st.divider()

# ============ Загрузка запросов КП ============
col_rb, col_fb = st.columns(2)

with col_rb:
    st.markdown('<div class="contract-header">🔵 РБ — Региональный бюджет</div>', unsafe_allow_html=True)
    rb_file = st.file_uploader(
        "Запрос КП (РБ)",
        type=['docx'],
        key="rb_uploader",
        help="Запрос на КП РБ (.docx)"
    )

    col_status, col_btn = st.columns([3, 1])
    with col_status:
        if st.session_state.rb_request is not None:
            st.success(f"✅ {len(st.session_state.rb_request)} позиций")
    with col_btn:
        if st.session_state.rb_request is not None:
            if st.button("🔄", key="reparse_rb", help="Перераспознать"):
                st.session_state.loaded_files['rb'] = None
                save_data('loaded_files', st.session_state.loaded_files)
                st.rerun()

    if rb_file:
        need_parse = (
            st.session_state.loaded_files['rb'] != rb_file.name or
            st.session_state.rb_request is None or
            len(st.session_state.rb_request) == 0
        )
        if need_parse:
            try:
                with st.spinner("📄 Парсинг Word..."):
                    st.session_state.rb_request = parse_request_file(rb_file, rb_file.name)
                if st.session_state.rb_request is not None and len(st.session_state.rb_request) > 0:
                    st.session_state.loaded_files['rb'] = rb_file.name
                    save_data('rb_request', st.session_state.rb_request)
                    save_data('loaded_files', st.session_state.loaded_files)
                else:
                    st.warning("⚠️ Не найдено ни одной позиции.")
                st.rerun()
            except Exception as e:
                st.error(f"Ошибка: {e}")

with col_fb:
    st.markdown('<div class="contract-header">🟢 ФБ — Федеральный бюджет</div>', unsafe_allow_html=True)
    fb_file = st.file_uploader(
        "Запрос КП (ФБ)",
        type=['docx'],
        key="fb_uploader",
        help="Запрос на КП ФБ (.docx)"
    )

    col_status, col_btn = st.columns([3, 1])
    with col_status:
        if st.session_state.fb_request is not None:
            st.success(f"✅ {len(st.session_state.fb_request)} позиций")
    with col_btn:
        if st.session_state.fb_request is not None:
            if st.button("🔄", key="reparse_fb", help="Перераспознать"):
                st.session_state.loaded_files['fb'] = None
                save_data('loaded_files', st.session_state.loaded_files)
                st.rerun()

    if fb_file:
        need_parse = (
            st.session_state.loaded_files['fb'] != fb_file.name or
            st.session_state.fb_request is None or
            len(st.session_state.fb_request) == 0
        )
        if need_parse:
            try:
                with st.spinner("📄 Парсинг Word..."):
                    st.session_state.fb_request = parse_request_file(fb_file, fb_file.name)
                if st.session_state.fb_request is not None and len(st.session_state.fb_request) > 0:
                    st.session_state.loaded_files['fb'] = fb_file.name
                    save_data('fb_request', st.session_state.fb_request)
                    save_data('loaded_files', st.session_state.loaded_files)
                else:
                    st.warning("⚠️ Не найдено ни одной позиции.")
                st.rerun()
            except Exception as e:
                st.error(f"Ошибка: {e}")

# ============ Кнопка расчёта ============
st.divider()
if st.button("🧮 Рассчитать КП", type="primary", use_container_width=True):
    if st.session_state.cost_data is None:
        st.error("Загрузите файл себестоимости")
    elif st.session_state.competitor_data is None:
        st.error("Загрузите КП конкурента")
    elif st.session_state.rb_request is None and st.session_state.fb_request is None:
        st.error("Загрузите хотя бы один запрос КП (РБ или ФБ)")
    else:
        with st.spinner("Расчёт..."):
            # РБ
            if st.session_state.rb_request is not None:
                matched_rb = match_products(
                    st.session_state.rb_request,
                    st.session_state.cost_data,
                    st.session_state.competitor_data
                )
                priced_rb = calculate_prices(matched_rb)
                priced_rb['Контракт'] = 'РБ'
                priced_rb['№'] = range(1, len(priced_rb) + 1)
                st.session_state.rb_data = priced_rb
                save_data('rb_data', priced_rb)

            # ФБ
            if st.session_state.fb_request is not None:
                matched_fb = match_products(
                    st.session_state.fb_request,
                    st.session_state.cost_data,
                    st.session_state.competitor_data
                )
                priced_fb = calculate_prices(matched_fb)
                priced_fb['Контракт'] = 'ФБ'
                priced_fb['№'] = range(1, len(priced_fb) + 1)
                st.session_state.fb_data = priced_fb
                save_data('fb_data', priced_fb)

            # Очистить старые ключи editor-ов
            for key in ['rb_editor', 'fb_editor']:
                if key in st.session_state:
                    del st.session_state[key]

        st.success("✅ Расчёт выполнен!")
        st.rerun()


# ============ ФУНКЦИИ ДЛЯ ОТРИСОВКИ ============

def render_mini_dashboard(df, container):
    """Компактный дашборд для одного контракта"""
    econ = calculate_economics(df)
    margin_color = "green" if econ['margin_percent'] >= 0 else "red"
    loss_color = "red" if econ['loss_positions'] > 0 else "green"

    container.markdown(f"""
    <div class="metrics-row">
        <div class="metric-card">
            <div class="metric-label">Позиций</div>
            <div class="metric-value gray">{econ['total_positions']} (конк. {econ['positions_with_comp']})</div>
        </div>
        <div class="metric-card">
            <div class="metric-label">Сумма контракта</div>
            <div class="metric-value teal">{econ['contract_total']:,.0f} ₽</div>
        </div>
        <div class="metric-card">
            <div class="metric-label">Себестоимость</div>
            <div class="metric-value gray">{econ['cost_total']:,.0f} ₽</div>
        </div>
        <div class="metric-card">
            <div class="metric-label">Маржа</div>
            <div class="metric-value {margin_color}">{econ['profit']:,.0f} ₽ ({econ['margin_percent']:.1f}%)</div>
        </div>
        <div class="metric-card">
            <div class="metric-label">Скидка от конкурента</div>
            <div class="metric-value orange">{econ['discount_percent']:.2f}%</div>
        </div>
        <div class="metric-card">
            <div class="metric-label">Убыточных</div>
            <div class="metric-value {loss_color}">{econ['loss_positions']} шт / {econ['loss_total_rub']:,.0f} ₽</div>
        </div>
    </div>
    """, unsafe_allow_html=True)


def render_contract_section(state_key, label, emoji, header_class, editor_key):
    """Секция контракта: дашборд + таблица + управление"""
    df = st.session_state.get(state_key)
    if df is None:
        return

    st.markdown(f'<div class="{header_class}">{emoji} {label}</div>', unsafe_allow_html=True)

    # Плейсхолдер для дашборда (заполним после таблицы, чтобы данные были актуальные)
    dashboard_placeholder = st.container()

    # Подготовка таблицы
    display_df = df.copy()
    display_df['Маржа %'] = ((display_df['Наша цена'] - display_df['Себестоимость']) / display_df['Наша цена'] * 100).replace([float('inf'), float('-inf')], 0).fillna(0).round(1)
    display_df['Маржа руб'] = ((display_df['Наша цена'] - display_df['Себестоимость']) * display_df['Кол-во']).round(2)

    if 'Описание' not in display_df.columns:
        display_df['Описание'] = ''

    edit_columns = ['Наименование', 'Описание', 'Ед.изм.', 'Кол-во', 'Себестоимость',
                    'Цена конкурента', 'Наша цена', 'Маржа %', 'Маржа руб', 'Тара', 'Матч']
    edit_columns = [c for c in edit_columns if c in display_df.columns]
    show_df = display_df[edit_columns].copy()

    edited = st.data_editor(
        show_df,
        num_rows="fixed",
        use_container_width=True,
        key=editor_key,
        height=500,
        disabled=['Маржа %', 'Маржа руб', 'Тара', 'Матч'],
        column_config={
            "Наименование": st.column_config.TextColumn("Наименование", width="large"),
            "Описание": st.column_config.TextColumn("Описание", width="large"),
            "Ед.изм.": st.column_config.TextColumn("Ед.изм.", width="small"),
            "Кол-во": st.column_config.NumberColumn("Кол-во", format="%.1f", step=1, width="small"),
            "Себестоимость": st.column_config.NumberColumn("Себес", format="%.2f", step=0.5, width="small"),
            "Цена конкурента": st.column_config.NumberColumn("Конкурент", format="%.2f", step=0.5, width="small"),
            "Наша цена": st.column_config.NumberColumn("Наша цена", format="%.2f", step=0.5, width="small"),
            "Маржа %": st.column_config.NumberColumn("Маржа %", format="%.1f", width="small"),
            "Маржа руб": st.column_config.NumberColumn("Маржа руб", format="%.0f", width="small"),
            "Тара": st.column_config.TextColumn("Тара", width="large"),
            "Матч": st.column_config.TextColumn("Матч", width="large"),
        }
    )

    # Синхронизируем правки в state
    for col in ['Наименование', 'Описание', 'Ед.изм.', 'Кол-во', 'Себестоимость', 'Наша цена', 'Цена конкурента']:
        if col in edited.columns:
            st.session_state[state_key][col] = edited[col]
    save_data(state_key, st.session_state[state_key])

    # Заполняем дашборд (теперь с актуальными данными после синхронизации)
    render_mini_dashboard(st.session_state[state_key], dashboard_placeholder)

    # Управление: наценка + пересчитать + скачать таблицу
    markup_key = f'markup_{state_key}'
    if markup_key not in st.session_state:
        st.session_state[markup_key] = 30.0

    col_markup, col_recalc, col_download = st.columns([2, 2, 1])
    with col_markup:
        st.session_state[markup_key] = st.number_input(
            "Наценка без конкурента (%)",
            min_value=0.0, max_value=200.0, value=st.session_state[markup_key], step=1.0,
            key=f"markup_input_{state_key}"
        )
    with col_recalc:
        if st.button("🔄 Пересчитать", key=f"recalc_{state_key}", use_container_width=True):
            markup = st.session_state[markup_key] / 100
            data = st.session_state[state_key]
            for idx in data.index:
                comp = float(data.at[idx, 'Цена конкурента'] or 0)
                cost = float(data.at[idx, 'Себестоимость'] or 0)
                if comp <= 0 and cost > 0:
                    data.at[idx, 'Наша цена'] = round(cost * (1 + markup), 2)
            data['Сумма'] = (data['Наша цена'] * data['Кол-во']).round(2)
            data['Маржа'] = (data['Наша цена'] - data['Себестоимость']).round(2)
            data['Маржа %'] = (data['Маржа'] / data['Наша цена'] * 100).replace([float('inf'), float('-inf')], 0).fillna(0).round(1)
            # Очистить состояние editor-а чтобы старые правки не наложились
            if editor_key in st.session_state:
                del st.session_state[editor_key]
            save_data(state_key, data)
            st.rerun()
    with col_download:
        export_cols = ['Наименование', 'Ед.изм.', 'Кол-во', 'Себестоимость',
                       'Цена конкурента', 'Наша цена']
        export_cols = [c for c in export_cols if c in st.session_state[state_key].columns]
        buf = BytesIO()
        export_df = st.session_state[state_key][export_cols].copy()
        export_df['Маржа %'] = ((st.session_state[state_key]['Наша цена'] - st.session_state[state_key]['Себестоимость']) / st.session_state[state_key]['Наша цена'] * 100).replace([float('inf'), float('-inf')], 0).fillna(0).round(1)
        export_df['Маржа руб'] = ((st.session_state[state_key]['Наша цена'] - st.session_state[state_key]['Себестоимость']) * st.session_state[state_key]['Кол-во']).round(2)
        export_df.to_excel(buf, index=False, sheet_name=label)
        st.download_button(
            "📥 Excel",
            data=buf.getvalue(),
            file_name=f"Таблица_{label}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
            key=f"dl_table_{state_key}"
        )


# ============ ШАГ 2: Таблицы и дашборды ============

has_rb = st.session_state.get('rb_data') is not None
has_fb = st.session_state.get('fb_data') is not None

if has_rb or has_fb:
    st.divider()
    st.subheader("📊 Сводка")

    # Плейсхолдер для общей сводки (заполним ПОСЛЕ секций контрактов, чтобы данные были актуальные)
    summary_placeholder = st.container()

    st.caption("Редактируйте данные в таблицах ниже. Дашборд обновляется автоматически.")

    # Секции контрактов (синхронизируют правки из data_editor)
    if has_rb:
        st.divider()
        render_contract_section('rb_data', 'РБ', '🔵', 'contract-header-rb', 'rb_editor')

    if has_fb:
        st.divider()
        render_contract_section('fb_data', 'ФБ', '🟢', 'contract-header-fb', 'fb_editor')

    # Заполняем сводку ПОСЛЕ синхронизации правок
    with summary_placeholder:
        if has_rb and has_fb:
            rb_econ = calculate_economics(st.session_state.rb_data)
            fb_econ = calculate_economics(st.session_state.fb_data)
            combined = pd.concat([st.session_state.rb_data, st.session_state.fb_data], ignore_index=True)
            total_econ = calculate_economics(combined)

            st.markdown(f"""
            <div class="summary-row">
                <div class="summary-card rb">
                    <div class="summary-title">🔵 РБ — {rb_econ['total_positions']} поз.</div>
                    <div class="summary-line">
                        Сумма: <b>{rb_econ['contract_total']:,.0f} ₽</b> &nbsp;|&nbsp;
                        Маржа: <b>{rb_econ['profit']:,.0f} ₽</b> ({rb_econ['margin_percent']:.1f}%) &nbsp;|&nbsp;
                        Скидка: <b>{rb_econ['discount_percent']:.2f}%</b> &nbsp;|&nbsp;
                        Убыт.: <b>{rb_econ['loss_positions']}</b>
                    </div>
                </div>
                <div class="summary-card fb">
                    <div class="summary-title">🟢 ФБ — {fb_econ['total_positions']} поз.</div>
                    <div class="summary-line">
                        Сумма: <b>{fb_econ['contract_total']:,.0f} ₽</b> &nbsp;|&nbsp;
                        Маржа: <b>{fb_econ['profit']:,.0f} ₽</b> ({fb_econ['margin_percent']:.1f}%) &nbsp;|&nbsp;
                        Скидка: <b>{fb_econ['discount_percent']:.2f}%</b> &nbsp;|&nbsp;
                        Убыт.: <b>{fb_econ['loss_positions']}</b>
                    </div>
                </div>
                <div class="summary-card total">
                    <div class="summary-title">ИТОГО — {total_econ['total_positions']} поз.</div>
                    <div class="summary-line">
                        Сумма: <b>{total_econ['contract_total']:,.0f} ₽</b> &nbsp;|&nbsp;
                        Себес: <b>{total_econ['cost_total']:,.0f} ₽</b> &nbsp;|&nbsp;
                        Маржа: <b>{total_econ['profit']:,.0f} ₽</b> ({total_econ['margin_percent']:.1f}%) &nbsp;|&nbsp;
                        Убыт.: <b>{total_econ['loss_positions']}</b>
                    </div>
                </div>
            </div>
            """, unsafe_allow_html=True)
        elif has_rb:
            rb_econ = calculate_economics(st.session_state.rb_data)
            st.markdown(f"""
            <div class="summary-row">
                <div class="summary-card rb">
                    <div class="summary-title">🔵 РБ — {rb_econ['total_positions']} поз.</div>
                    <div class="summary-line">
                        Сумма: <b>{rb_econ['contract_total']:,.0f} ₽</b> &nbsp;|&nbsp;
                        Маржа: <b>{rb_econ['profit']:,.0f} ₽</b> ({rb_econ['margin_percent']:.1f}%) &nbsp;|&nbsp;
                        Скидка: <b>{rb_econ['discount_percent']:.2f}%</b>
                    </div>
                </div>
            </div>
            """, unsafe_allow_html=True)
        elif has_fb:
            fb_econ = calculate_economics(st.session_state.fb_data)
            st.markdown(f"""
            <div class="summary-row">
                <div class="summary-card fb">
                    <div class="summary-title">🟢 ФБ — {fb_econ['total_positions']} поз.</div>
                    <div class="summary-line">
                        Сумма: <b>{fb_econ['contract_total']:,.0f} ₽</b> &nbsp;|&nbsp;
                        Маржа: <b>{fb_econ['profit']:,.0f} ₽</b> ({fb_econ['margin_percent']:.1f}%) &nbsp;|&nbsp;
                        Скидка: <b>{fb_econ['discount_percent']:.2f}%</b>
                    </div>
                </div>
            </div>
            """, unsafe_allow_html=True)

    # ============ ШАГ 3: Скачать документы ============
    st.divider()
    st.subheader("📥 Шаг 3: Скачать документы")

    download_cols = st.columns(4)
    col_idx = 0

    if has_rb:
        rb_df = st.session_state.rb_data.copy()
        rb_df['№'] = range(1, len(rb_df) + 1)
        with download_cols[col_idx]:
            docx_rb = export_kp_to_docx(rb_df, "РБ")
            st.download_button(
                "📄 КП_РБ.docx",
                data=docx_rb,
                file_name="КП_РБ.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        col_idx += 1
        with download_cols[col_idx]:
            excel_rb = export_kp_to_excel(rb_df, "РБ")
            st.download_button(
                "📥 Excel РБ",
                data=excel_rb,
                file_name="КП_РБ.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
        col_idx += 1

    if has_fb:
        fb_df = st.session_state.fb_data.copy()
        fb_df['№'] = range(1, len(fb_df) + 1)
        with download_cols[col_idx]:
            docx_fb = export_kp_to_docx(fb_df, "ФБ")
            st.download_button(
                "📄 КП_ФБ.docx",
                data=docx_fb,
                file_name="КП_ФБ.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        col_idx += 1
        with download_cols[col_idx]:
            excel_fb = export_kp_to_excel(fb_df, "ФБ")
            st.download_button(
                "📥 Excel ФБ",
                data=excel_fb,
                file_name="КП_ФБ.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )

# ============ Секретный инструмент: PDF → Word (в сайдбаре) ============
with st.sidebar:
    st.markdown("#### PDF → Word")
    st.caption("Извлечение таблиц из PDF через Claude Opus")

    _api_key = os.getenv("OPENROUTER_API_KEY", "")
    _pdf_file = st.file_uploader("PDF файл", type=['pdf'], key="secret_pdf")

    _process = st.button("Обработать", key="process_pdf_btn", use_container_width=True)

    if _process:
        if not _api_key:
            st.error("Введите API ключ")
        elif not _pdf_file:
            st.error("Загрузите PDF")
        else:
            with st.spinner("Claude Opus анализирует PDF..."):
                _error_msg = None
                _raw_content = None
                try:
                    _pdf_bytes = _pdf_file.read()
                    _pdf_b64 = base64.b64encode(_pdf_bytes).decode()

                    _prompt = (
                        "Извлеки ВСЕ таблицы из этого PDF документа. "
                        "Верни результат СТРОГО в JSON формате без markdown обёрток:\n"
                        '{"tables": [{"name": "Название таблицы или пустая строка", '
                        '"headers": ["Столбец1", "Столбец2"], '
                        '"rows": [["значение1", "значение2"]]}]}\n\n'
                        "ВАЖНО:\n"
                        "- Сохрани ВСЕ данные точно как в оригинале\n"
                        "- Числа оставь как строки\n"
                        "- Объединённые ячейки раздели на отдельные\n"
                        "- Порядок столбцов и строк точно как в оригинале\n"
                        "- Все таблицы из документа\n"
                        "- НЕ оборачивай в ```json``` — чистый JSON"
                    )

                    _resp = http_requests.post(
                        'https://openrouter.ai/api/v1/chat/completions',
                        headers={
                            'Authorization': f'Bearer {_api_key}',
                            'Content-Type': 'application/json',
                            'HTTP-Referer': 'https://krechet.space',
                        },
                        json={
                            'model': 'anthropic/claude-opus-4-6',
                            'messages': [{
                                'role': 'user',
                                'content': [
                                    {
                                        'type': 'file',
                                        'file': {
                                            'filename': _pdf_file.name,
                                            'content': _pdf_b64
                                        }
                                    },
                                    {
                                        'type': 'text',
                                        'text': _prompt
                                    }
                                ]
                            }],
                            'max_tokens': 32000,
                            'temperature': 0,
                        },
                        timeout=180
                    )

                    if _resp.status_code != 200:
                        _error_msg = f"Ошибка API ({_resp.status_code})"
                        _raw_content = _resp.text[:1500]
                    else:
                        _data = _resp.json()
                        _raw_content = _data.get('choices', [{}])[0].get('message', {}).get('content', '')

                        if not _raw_content:
                            _error_msg = "Пустой ответ от AI"
                        else:
                            _cleaned = _raw_content.strip()
                            if _cleaned.startswith('```'):
                                _cleaned = _cleaned.split('\n', 1)[-1]
                            if _cleaned.endswith('```'):
                                _cleaned = _cleaned.rsplit('```', 1)[0]
                            _cleaned = _cleaned.strip()

                            _tables_data = json.loads(_cleaned)

                            if not _tables_data.get('tables'):
                                _error_msg = "Таблицы не найдены в документе"
                            else:
                                from docx import Document as DocxDoc
                                from docx.shared import Pt

                                doc = DocxDoc()
                                for tbl in _tables_data['tables']:
                                    if tbl.get('name'):
                                        doc.add_heading(tbl['name'], level=2)

                                    headers = tbl.get('headers', [])
                                    rows = tbl.get('rows', [])
                                    col_count = max(
                                        len(headers),
                                        max((len(r) for r in rows), default=0)
                                    )

                                    if col_count == 0:
                                        continue

                                    row_count = len(rows) + (1 if headers else 0)
                                    t = doc.add_table(rows=row_count, cols=col_count)
                                    t.style = 'Table Grid'

                                    if headers:
                                        for i, h in enumerate(headers):
                                            if i < col_count:
                                                cell = t.rows[0].cells[i]
                                                cell.text = str(h)
                                                for p in cell.paragraphs:
                                                    for run in p.runs:
                                                        run.bold = True
                                                        run.font.size = Pt(10)

                                    start_row = 1 if headers else 0
                                    for ri, row in enumerate(rows):
                                        for ci, val in enumerate(row):
                                            if ci < col_count:
                                                cell = t.rows[start_row + ri].cells[ci]
                                                cell.text = str(val or '')
                                                for p in cell.paragraphs:
                                                    for run in p.runs:
                                                        run.font.size = Pt(10)

                                    doc.add_paragraph()

                                _buf = BytesIO()
                                doc.save(_buf)
                                st.session_state._pdf_docx = _buf.getvalue()
                                st.session_state._pdf_name = _pdf_file.name.replace('.pdf', '.docx')
                                st.success(f"Найдено таблиц: {len(_tables_data['tables'])}")

                except json.JSONDecodeError:
                    _error_msg = "Не удалось распознать ответ AI как JSON"
                except Exception as e:
                    _error_msg = str(e)

                if _error_msg:
                    st.error(_error_msg)
                    if _raw_content:
                        with st.expander("Ответ AI"):
                            st.code(_raw_content[:3000])

    if st.session_state.get('_pdf_docx'):
        st.download_button(
            "Скачать Word",
            data=st.session_state._pdf_docx,
            file_name=st.session_state.get('_pdf_name', 'result.docx'),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="download_pdf_docx"
        )

# Подвал
st.divider()
col_footer, col_clear = st.columns([4, 1])
with col_footer:
    st.caption("© 2026 Сервис расчёта КП | v4.0 | Данные сохраняются автоматически")
with col_clear:
    if st.button("🗑️ Очистить всё", type="secondary"):
        for f in CACHE_DIR.glob("*.pkl"):
            f.unlink()
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()
