"""
Экспорт КП в Word (.docx)
Формат по образцу КП конкурента: таблица с ценами
"""

import math
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def format_number(value: float, decimals: int = 2) -> str:
    """Форматирование числа в русском стиле: 1 234,56"""
    if value is None or (isinstance(value, float) and (math.isnan(value) or math.isinf(value))):
        return "0"
    if value == 0:
        return "0"
    if decimals == 0:
        formatted = f"{value:,.0f}"
    else:
        formatted = f"{value:,.{decimals}f}"
    # Меняем разделители на русские
    formatted = formatted.replace(",", " ").replace(".", ",")
    return formatted


def safe_str(value) -> str:
    """Безопасное преобразование в строку: NaN, None, pd.NA → пустая строка"""
    if value is None:
        return ''
    if isinstance(value, float) and (math.isnan(value) or math.isinf(value)):
        return ''
    s = str(value)
    if s in ('nan', 'None', '<NA>', 'NaT'):
        return ''
    return s


def safe_float(value, default=0) -> float:
    """Безопасное преобразование в float: NaN, None → default"""
    if value is None:
        return default
    if isinstance(value, (int, float)):
        if math.isnan(value) or math.isinf(value):
            return default
        return float(value)
    try:
        return float(value)
    except (ValueError, TypeError):
        return default


def set_cell_text(cell, text: str, bold: bool = False, align: str = 'left', size: int = 9):
    """Установить текст ячейки с форматированием"""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(safe_str(text))
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.bold = bold

    if align == 'center':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Уменьшаем отступы внутри ячеек
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(1)


def set_cell_shading(cell, color: str):
    """Установить цвет фона ячейки"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color
    })
    shading.append(shading_elem)


def export_kp_to_docx(df: pd.DataFrame, contract_type: str = "КП") -> bytes:
    """
    Экспорт КП в Word документ по формату конкурента.

    Args:
        df: DataFrame с данными КП (Наименование, Описание, Ед.изм., Кол-во, Наша цена)
        contract_type: Тип контракта (РБ/ФБ)

    Returns:
        Байты .docx файла
    """
    doc = Document()

    # Настройки страницы — портретная ориентация
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Колонки таблицы
    headers = ['№ п/п', 'Наименование продукта', 'Описание / ГОСТ', 'Ед. изм.',
               'Количество', 'Цена за ед., руб.', 'Общая стоимость, руб.']

    num_rows = len(df) + 2  # заголовок + данные + итого
    table = doc.add_table(rows=num_rows, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Применяем стиль таблицы с рамками
    table.style = 'Table Grid'

    # Заголовки
    header_color = '4472C4'
    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)
        set_cell_text(cell, header, bold=True, align='center', size=9)
        set_cell_shading(cell, header_color)
        # Белый текст на синем фоне
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Данные
    total_sum = 0
    for row_num, (_, row) in enumerate(df.iterrows(), start=1):
        price = safe_float(row.get('Наша цена', 0))
        qty = safe_float(row.get('Кол-во', 0))
        sum_value = round(price * qty, 2)
        total_sum += sum_value

        # № п/п
        set_cell_text(table.cell(row_num, 0), str(row_num), align='center')
        # Наименование
        set_cell_text(table.cell(row_num, 1), str(row.get('Наименование', '')), align='left')
        # Описание / ГОСТ
        set_cell_text(table.cell(row_num, 2), str(row.get('Описание', '') or ''), align='left')
        # Ед. изм.
        set_cell_text(table.cell(row_num, 3), str(row.get('Ед.изм.', '')), align='center')
        # Количество
        set_cell_text(table.cell(row_num, 4), format_number(qty, 0), align='center')
        # Цена за ед.
        set_cell_text(table.cell(row_num, 5), format_number(price, 2), align='right')
        # Общая стоимость
        set_cell_text(table.cell(row_num, 6), format_number(sum_value, 2), align='right')

    # Строка ИТОГО
    last_row = len(df) + 1
    # Объединяем первые 6 колонок
    merge_cell = table.cell(last_row, 0)
    for col_idx in range(1, 6):
        merge_cell = merge_cell.merge(table.cell(last_row, col_idx))
    set_cell_text(merge_cell, 'ИТОГО:', bold=True, align='right', size=10)
    set_cell_text(table.cell(last_row, 6), format_number(total_sum, 2), bold=True, align='right', size=10)

    # Подсветка итоговой строки
    itogo_color = 'E2EFDA'
    for col_idx in range(len(headers)):
        try:
            set_cell_shading(table.cell(last_row, col_idx), itogo_color)
        except Exception:
            pass

    # Ширина колонок
    widths = [Cm(0.8), Cm(5.0), Cm(5.0), Cm(1.2), Cm(1.5), Cm(2.0), Cm(2.5)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = width

    # Сохраняем в байты
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()
